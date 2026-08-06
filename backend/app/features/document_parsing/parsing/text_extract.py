
from __future__ import annotations

import io
import logging
import os
import re
import tempfile
from functools import lru_cache
from pathlib import Path
from threading import Lock

from docx import Document

from app.core.errors import ApiError

logger = logging.getLogger(__name__)
_DOCLING_CONVERT_LOCK = Lock()
PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
# Minimum usable extracted text length (chars). Applied on every accept path so
# short/noise extracts rejected by the fast path cannot re-enter via a soft fallback.
MIN_PDF_TEXT_CHARS = 200
MIN_DOCX_TEXT_CHARS = 80


def _normalize_extracted_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00ad", "").replace("\ufeff", "")
    lines: list[str] = []
    for raw in text.split("\n"):
        line = re.sub(r"[ \t]+", " ", raw).strip()
        if line.startswith("#"):
            line = re.sub(r"^#+\s*", "", line).strip()
        lines.append(line)
    cleaned: list[str] = []
    blank_run = 0
    for line in lines:
        if not line:
            blank_run += 1
            if blank_run <= 1:
                cleaned.append("")
            continue
        blank_run = 0
        cleaned.append(line)
    return "\n".join(cleaned).strip()


def _require_docling():
    try:
        os.environ.setdefault("DOCLING_INFERENCE_COMPILE_TORCH_MODELS", "false")
        from docling.document_converter import DocumentConverter
        return DocumentConverter
    except Exception as exc:
        raise ApiError(
            503,
            "docling_not_installed",
            "Docling is required for accurate resume PDF parsing. "
            "Install it with: pip install 'docling>=2.0,<3'",
        ) from exc


@lru_cache(maxsize=1)
def _get_docling_converter():
    """Create Docling's heavyweight converter once per backend process."""
    return _require_docling()()


def _extract_pdf_fast(content: bytes) -> str | None:
    """Lightweight extractors before Docling. Returns None only when text is short.

    Raises ApiError for encrypted PDFs and for hard parse failures from the extractor.
    Unexpected exceptions are logged and re-raised as ApiError (not silently ignored).
    """
    try:
        from app.features.document_parsing.extractors.pdf import parse_pdf_to_blocks

        blocks = parse_pdf_to_blocks(content)
    except ApiError as exc:
        if exc.code == "encrypted_pdf":
            raise
        # Soft miss (e.g. empty/unsupported) → try Docling next.
        logger.info("pdf_fast_extract_skipped code=%s message=%s", exc.code, exc.message)
        return None
    except Exception as exc:
        logger.exception("pdf_fast_extract_unexpected")
        raise ApiError(
            400,
            "pdf_parse_failed",
            f"Lightweight PDF extraction failed: {type(exc).__name__}",
        ) from exc
    text = _normalize_extracted_text(
        "\n".join(block.text for block in blocks if getattr(block, "text", None))
    )
    return text if len(text) >= MIN_PDF_TEXT_CHARS else None


def _extract_docx_fast(content: bytes) -> str | None:
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        logger.info("docx_fast_extract_skipped type=%s", type(exc).__name__)
        return None
    text = _normalize_extracted_text("\n".join(_docx_paragraph_text(document)))
    return text if len(text) >= MIN_DOCX_TEXT_CHARS else None


def _extract_with_docling(content: bytes, suffix: str) -> str:
    tmp_path: str | None = None
    try:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as handle:
            handle.write(content)
            tmp_path = handle.name
        converter = _get_docling_converter()
        # Docling's converter is reused, but its heavyweight inference pipeline
        # is not safe to execute concurrently on the same CPU-bound instance.
        with _DOCLING_CONVERT_LOCK:
            result = converter.convert(tmp_path)
        doc = result.document
        text = ""
        if hasattr(doc, "export_to_markdown"):
            text = doc.export_to_markdown() or ""
        elif hasattr(doc, "export_to_text"):
            text = doc.export_to_text() or ""
        else:
            text = str(doc)
        text = _normalize_extracted_text(text)
        if not text:
            raise ApiError(
                422,
                "document_has_no_text",
                "Docling found no usable text in this document.",
            )
        return text
    except ApiError:
        raise
    except Exception as exc:
        logger.exception("docling_extract_failed")
        raise ApiError(
            400,
            "document_parse_failed",
            "Docling could not parse this document. Try a text-based PDF or DOCX.",
        ) from exc
    finally:
        if tmp_path:
            try:
                Path(tmp_path).unlink(missing_ok=True)
            except Exception:
                logger.warning("docling_temp_cleanup_failed path=%s", tmp_path)


def _docx_paragraph_text(document: Document) -> list[str]:
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            for cell in cells:
                for part in cell.splitlines():
                    part = part.strip()
                    if part:
                        lines.append(part)
    return lines


def _pdf_blocks_text(content: bytes) -> str:
    from app.features.document_parsing.extractors.pdf import parse_pdf_to_blocks

    blocks = parse_pdf_to_blocks(content)
    return _normalize_extracted_text(
        "\n".join(block.text for block in blocks if getattr(block, "text", None))
    )


def extract_text(content: bytes, mime_type: str) -> str:
    if not content:
        raise ApiError(400, "empty_document", "The selected document is empty.")
    if mime_type == PDF_MIME:
        fast_text = _extract_pdf_fast(content)
        if fast_text:
            return fast_text
        try:
            docling_text = _extract_with_docling(content, ".pdf")
        except ApiError as exc:
            # Fail closed when Docling finds no usable text — do not accept short
            # garbage from secondary extractors (fail-open anti-pattern).
            if exc.code == "document_has_no_text":
                raise
            # Only when Docling is missing may lightweight extractors serve PDF
            # text, and only if they meet the quality floor.
            if exc.code == "docling_not_installed":
                try:
                    text = _pdf_blocks_text(content)
                except ApiError:
                    raise
                if len(text) >= MIN_PDF_TEXT_CHARS:
                    logger.warning("pdf_extract_used_lightweight_backend docling_missing=1")
                    return text
                raise ApiError(
                    422,
                    "document_has_no_text",
                    "No usable text was found in this PDF (Docling not installed).",
                ) from exc
            # document_parse_failed / other: allow one quality-gated secondary try
            if exc.code == "document_parse_failed":
                try:
                    text = _pdf_blocks_text(content)
                except ApiError as secondary:
                    raise ApiError(
                        400,
                        "document_parse_failed",
                        "Could not parse this PDF with Docling or secondary extractors.",
                    ) from secondary
                if len(text) >= MIN_PDF_TEXT_CHARS:
                    logger.warning(
                        "pdf_extract_used_lightweight_backend after_docling_parse_failed=1 chars=%s",
                        len(text),
                    )
                    return text
                raise ApiError(
                    422,
                    "document_has_no_text",
                    "No usable text was found in this PDF after Docling parse failure.",
                ) from exc
            raise
        if len(docling_text) < MIN_PDF_TEXT_CHARS:
            raise ApiError(
                422,
                "document_has_no_text",
                f"Extracted PDF text is too short to use (need at least {MIN_PDF_TEXT_CHARS} characters).",
            )
        return docling_text
    if mime_type == DOCX_MIME:
        fast_text = _extract_docx_fast(content)
        if fast_text:
            return fast_text
        try:
            docling_text = _extract_with_docling(content, ".docx")
        except ApiError as exc:
            if exc.code == "document_has_no_text":
                raise
            if exc.code in {"docling_not_installed", "document_parse_failed"}:
                try:
                    document = Document(io.BytesIO(content))
                except Exception as docx_exc:
                    raise ApiError(
                        400,
                        "document_parse_failed",
                        "Could not parse this DOCX document.",
                    ) from docx_exc
                text = _normalize_extracted_text("\n".join(_docx_paragraph_text(document)))
                if len(text) >= MIN_DOCX_TEXT_CHARS:
                    if exc.code == "document_parse_failed":
                        logger.warning("docx_extract_used_python_docx after_docling_parse_failed=1")
                    return text
                raise ApiError(
                    422,
                    "document_has_no_text",
                    "No usable text was found in the DOCX.",
                ) from exc
            raise
        if len(docling_text) < MIN_DOCX_TEXT_CHARS:
            raise ApiError(
                422,
                "document_has_no_text",
                f"Extracted DOCX text is too short to use (need at least {MIN_DOCX_TEXT_CHARS} characters).",
            )
        return docling_text
    raise ApiError(415, "unsupported_document_type", "Only PDF and DOCX documents are supported.")
